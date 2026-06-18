"""DOCX 解析器 — 段落 + 表格 + 标题"""
import logging
from typing import List
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from src.parser.base import BaseParser, ParseResult
logger = logging.getLogger("course_assistant.parser.docx")

class DocxParser(BaseParser):
    @property
    def supported_extensions(self) -> List[str]: return [".docx", ".DOCX"]

    def parse(self, file_path):
        result = ParseResult()
        try:
            doc = Document(file_path)
            props = doc.core_properties
            result.metadata = {"title": props.title or "", "author": props.author or ""}
            paragraphs, headings = [], []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text: continue
                if para.style and para.style.type == WD_STYLE_TYPE.PARAGRAPH:
                    sn = para.style.name
                    if "Heading 1" in sn: headings.append({"level":1,"text":text})
                    elif "Heading 2" in sn: headings.append({"level":2,"text":text})
                paragraphs.append(text)
            result.text = "\n".join(paragraphs); result.metadata["headings"] = headings
            tables = []
            for t in doc.tables:
                rows = [[c.text.strip() for c in r.cells] for r in t.rows]
                if rows: tables.append(self._table_to_markdown(rows))
            result.tables = tables
            if tables: result.text += "\n\n## 表格\n\n" + "\n\n".join(tables)
        except Exception as e: result.errors.append(f"DOCX error: {e}")
        return result
